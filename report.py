# -*- coding: utf-8 -*-
from __future__ import annotations

from io import BytesIO
from typing import Dict, Iterable

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt

from src import (
    DatiPlatea,
    genera_note_platea,
    genera_verifiche_platea,
    tabella_input_platea,
    tabella_sintesi_platea,
)


def _fmt(value: object) -> str:
    if isinstance(value, float):
        return f"{value:.2f}"
    if isinstance(value, int):
        return str(value)
    return str(value)


def _add_table(doc: Document, df: pd.DataFrame, columns: Iterable[str] | None = None) -> None:
    view = df.copy()
    if columns is not None:
        view = view.loc[:, list(columns)]
    table = doc.add_table(rows=1, cols=len(view.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(view.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(view.columns):
            cells[idx].text = _fmt(row[col])
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def _pilastri_report(d: DatiPlatea) -> pd.DataFrame:
    cols = ["x", "y", "P_kN", "Mx_kNm", "My_kNm"]
    out = d.pilastri_df.copy()
    for col in cols:
        if col not in out.columns:
            out[col] = 0.0
    return out.loc[:, cols]


def _add_map_plot(doc: Document, data, title: str, cbar_label: str) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    values = data.to_numpy() if hasattr(data, "to_numpy") else data
    fig, ax = plt.subplots(figsize=(6.8, 4.2), dpi=160)
    try:
        im = ax.imshow(values, origin="lower", cmap="Greys")
        ax.set_title(title, loc="left", fontsize=10, fontweight="bold", color="#111111")
        ax.set_xlabel("Indice griglia x")
        ax.set_ylabel("Indice griglia y")
        fig.colorbar(im, ax=ax, label=cbar_label, shrink=0.85)
        fig.tight_layout()
        bio = BytesIO()
        fig.savefig(bio, format="png", dpi=160, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        bio.seek(0)
        doc.add_picture(bio, width=Cm(14.8))
    except Exception:
        plt.close(fig)
        raise


def create_word_report(
    dati_stat: DatiPlatea,
    risultati_stat: Dict,
    dati_sis: DatiPlatea,
    risultati_sis: Dict,
    q_amm: float,
    modello: str,
) -> bytes:
    doc = Document()
    doc.add_heading("Relazione tecnica PlateaFEM", 0)
    doc.add_paragraph(
        "Analisi di platea di fondazione su suolo alla Winkler con controlli automatici di pressione, "
        "cedimento e sintesi delle mappe di calcolo."
    )

    doc.add_heading("1. Dati di input", level=1)
    _add_table(doc, tabella_input_platea(dati_stat, q_amm))

    doc.add_heading("2. Carichi da pilastri", level=1)
    _add_table(doc, _pilastri_report(dati_stat))

    doc.add_heading("3. Sintesi risultati", level=1)
    _add_table(doc, tabella_sintesi_platea(dati_stat, risultati_stat, risultati_sis, q_amm))

    doc.add_heading("4. Verifiche", level=1)
    _add_table(doc, genera_verifiche_platea(dati_stat, risultati_stat, risultati_sis, q_amm))

    doc.add_heading("5. Grafici statici", level=1)
    _add_map_plot(doc, risultati_stat["pressioni_kPa"], "Pressioni di contatto - caso statico", "kPa")
    _add_map_plot(doc, risultati_sis["pressioni_kPa"], "Pressioni di contatto - caso sismico", "kPa")
    _add_map_plot(doc, risultati_stat["cedimenti_mm"], "Cedimenti - caso statico", "mm")

    doc.add_heading("6. Note tecniche", level=1)
    for note in genera_note_platea(modello):
        doc.add_paragraph(note, style="List Bullet")

    doc.add_paragraph(
        "Documento generato automaticamente. I risultati devono essere verificati dal progettista responsabile "
        "in funzione del modello geotecnico e delle verifiche strutturali di dettaglio."
    )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
